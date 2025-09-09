from typing import Dict, List, Optional
import jinja2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
import base64
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self):
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader('templates/')
        )
        self.styles = getSampleStyleSheet()
        
    async def generate_documents(
        self, project_data: Dict, templates: List[str]
    ) -> Dict[str, str]:
        """Gera documentos do projeto em diferentes formatos"""
        documents = {}
        
        for template_name in templates:
            if template_name == "proposal":
                documents["proposal_docx"] = await self.generate_proposal_docx(project_data)
                documents["proposal_pdf"] = await self.generate_proposal_pdf(project_data)
            elif template_name == "budget":
                documents["budget_xlsx"] = await self.generate_budget_excel(project_data)
            elif template_name == "workplan":
                documents["workplan_pdf"] = await self.generate_workplan_pdf(project_data)
        
        return documents
    
    async def generate_proposal_docx(self, project_data: Dict) -> str:
        """Gera proposta em formato DOCX"""
        document = Document()
        
        # Configurar estilos
        self._setup_document_styles(document)
        
        # Capa
        self._add_cover_page(document, project_data)
        
        # Sumário
        document.add_page_break()
        self._add_table_of_contents(document)
        
        # Seções do projeto
        document.add_page_break()
        self._add_section(document, "1. IDENTIFICAÇÃO DO PROJETO", level=1)
        self._add_project_identification(document, project_data)
        
        self._add_section(document, "2. JUSTIFICATIVA", level=1)
        document.add_paragraph(project_data.get('justification', ''))
        
        self._add_section(document, "3. OBJETIVOS", level=1)
        self._add_objectives(document, project_data.get('objectives', {}))
        
        self._add_section(document, "4. METODOLOGIA", level=1)
        self._add_methodology(document, project_data.get('methodology', {}))
        
        self._add_section(document, "5. CRONOGRAMA", level=1)
        self._add_timeline(document, project_data.get('timeline', []))
        
        self._add_section(document, "6. ORÇAMENTO", level=1)
        self._add_budget(document, project_data.get('budget', {}))
        
        self._add_section(document, "7. RESULTADOS ESPERADOS", level=1)
        self._add_expected_results(document, project_data.get('expected_results', []))
        
        self._add_section(document, "8. EQUIPE", level=1)
        self._add_team(document, project_data.get('team', []))
        
        self._add_section(document, "9. INDICADORES DE AVALIAÇÃO", level=1)
        self._add_metrics(document, project_data.get('evaluation_metrics', []))
        
        self._add_section(document, "10. SUSTENTABILIDADE", level=1)
        self._add_sustainability(document, project_data.get('sustainability', {}))
        
        self._add_section(document, "11. ANÁLISE DE RISCOS", level=1)
        self._add_risks(document, project_data.get('risks', []))
        
        # Salvar documento
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        # Retornar como base64
        return base64.b64encode(buffer.read()).decode('utf-8')
    
    def _setup_document_styles(self, document):
        """Configura estilos do documento"""
        # Definir margens
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    
    def _add_cover_page(self, document, project_data: Dict):
        """Adiciona página de capa"""
        # Logo (se disponível)
        # document.add_picture('logo.png', width=Inches(2))
        
        # Título
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("PROGRAMA NACIONAL DE APOIO À ATENÇÃO DA SAÚDE\n")
        run.font.size = Pt(16)
        run.font.bold = True
        
        run = title.add_run("DA PESSOA COM DEFICIÊNCIA\n")
        run.font.size = Pt(16)
        run.font.bold = True
        
        run = title.add_run("PRONAS/PCD\n\n")
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Título do projeto
        project_title = document.add_paragraph()
        project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project_title.add_run(project_data.get('title', 'PROJETO'))
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 139)
        
        # Espaçamento
        for _ in range(10):
            document.add_paragraph()
        
        # Instituição
        institution = document.add_paragraph()
        institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = institution.add_run(f"Instituição Proponente\n")
        run.font.size = Pt(12)
        
        # Data
        date_para = document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(datetime.now().strftime("%B de %Y"))
        run.font.size = Pt(12)
    
    def _add_table_of_contents(self, document):
        """Adiciona sumário"""
        document.add_heading('SUMÁRIO', level=1)
        
        toc_items = [
            "1. IDENTIFICAÇÃO DO PROJETO",
            "2. JUSTIFICATIVA",
            "3. OBJETIVOS",
            "4. METODOLOGIA",
            "5. CRONOGRAMA",
            "6. ORÇAMENTO",
            "7. RESULTADOS ESPERADOS",
            "8. EQUIPE",
            "9. INDICADORES DE AVALIAÇÃO",
            "10. SUSTENTABILIDADE",
            "11. ANÁLISE DE RISCOS"
        ]
        
        for item in toc_items:
            para = document.add_paragraph()
            para.add_run(item)
            para.paragraph_format.left_indent = Inches(0.5)
    
    def _add_section(self, document, title: str, level: int = 1):
        """Adiciona uma seção ao documento"""
        document.add_heading(title, level=level)
    
    def _add_project_identification(self, document, project_data: Dict):
        """Adiciona identificação do projeto"""
        table = document.add_table(rows=5, cols=2)
        table.style = 'Light Grid'
        
        data = [
            ("Título:", project_data.get('title', '')),
            ("Instituição:", project_data.get('institution_name', '')),
            ("CNPJ:", project_data.get('institution_cnpj', '')),
            ("Tipo de Projeto:", project_data.get('type', '')),
            ("Duração:", f"{len(project_data.get('timeline', []))*3} meses")
        ]
        
        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    def _add_objectives(self, document, objectives: Dict):
        """Adiciona objetivos ao documento"""
        document.add_heading('Objetivo Geral', level=2)
        document.add_paragraph(objectives.get('general', ''))
        
        document.add_heading('Objetivos Específicos', level=2)
        for obj in objectives.get('specific', []):
            document.add_paragraph(f"• {obj}", style='List Bullet')
    
    def _add_methodology(self, document, methodology: Dict):
        """Adiciona metodologia ao documento"""
        document.add_heading('Abordagem Metodológica', level=2)
        document.add_paragraph(methodology.get('approach', ''))
        
        document.add_heading('Fases do Projeto', level=2)
        for phase in methodology.get('phases', []):
            document.add_paragraph(f"• {phase}", style='List Bullet')
        
        document.add_heading('Técnicas e Ferramentas', level=2)
        for technique in methodology.get('techniques', []):
            document.add_paragraph(f"• {technique}", style='List Bullet')
    
    def _add_timeline(self, document, timeline: List[Dict]):
        """Adiciona cronograma ao documento"""
        if not timeline:
            return
        
        table = document.add_table(rows=len(timeline)+1, cols=4)
        table.style = 'Light Grid'
        
        # Cabeçalho
        headers = ['Fase', 'Início', 'Fim', 'Entregas']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Dados
        for i, phase in enumerate(timeline):
            table.cell(i+1, 0).text = phase.get('phase', '')
            table.cell(i+1, 1).text = f"Mês {phase.get('start_month', '')}"
            table.cell(i+1, 2).text = f"Mês {phase.get('end_month', '')}"
            table.cell(i+1, 3).text = ', '.join(phase.get('deliverables', []))
    
    def _add_budget(self, document, budget: Dict):
        """Adiciona orçamento ao documento"""
        document.add_heading('Resumo Orçamentário', level=2)
        
        total = budget.get('total', 0)
        document.add_paragraph(f"Valor Total do Projeto: R$ {total:,.2f}")
        
        document.add_heading('Distribuição por Categoria', level=2)
        
        if budget.get('distribution'):
            table = document.add_table(rows=len(budget['distribution'])+1, cols=2)
            table.style = 'Light Grid'
            
            table.cell(0, 0).text = 'Categoria'
            table.cell(0, 1).text = 'Valor (R$)'
            
            for i, (category, value) in enumerate(budget['distribution'].items()):
                table.cell(i+1, 0).text = category.replace('_', ' ').title()
                table.cell(i+1, 1).text = f"{value:,.2f}"
    
    def _add_expected_results(self, document, results: List[str]):
        """Adiciona resultados esperados"""
        for result in results:
            document.add_paragraph(f"• {result}", style='List Bullet')
    
    def _add_team(self, document, team: List[Dict]):
        """Adiciona equipe ao documento"""
        if not team:
            return
        
        table = document.add_table(rows=len(team)+1, cols=4)
        table.style = 'Light Grid'
        
        headers = ['Função', 'Quantidade', 'Horas/Semana', 'Qualificação']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, member in enumerate(team):
            table.cell(i+1, 0).text = member.get('role', '')
            table.cell(i+1, 1).text = str(member.get('quantity', ''))
            table.cell(i+1, 2).text = str(member.get('hours_per_week', ''))
            table.cell(i+1, 3).text = member.get('qualifications', '')
    
    def _add_metrics(self, document, metrics: List[Dict]):
        """Adiciona métricas de avaliação"""
        if not metrics:
            return
        
        table = document.add_table(rows=len(metrics)+1, cols=4)
        table.style = 'Light Grid'
        
        headers = ['Indicador', 'Meta', 'Forma de Medição', 'Frequência']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, metric in enumerate(metrics):
            table.cell(i+1, 0).text = metric.get('indicator', '')
            table.cell(i+1, 1).text = str(metric.get('target', ''))
            table.cell(i+1, 2).text = metric.get('measurement', '')
            table.cell(i+1, 3).text = metric.get('frequency', '')
    
    def _add_sustainability(self, document, sustainability: Dict):
        """Adiciona plano de sustentabilidade"""
        for category, items in sustainability.items():
            document.add_heading(category.replace('_', ' ').title(), level=2)
            for item in items:
                document.add_paragraph(f"• {item}", style='List Bullet')
    
    def _add_risks(self, document, risks: List[Dict]):
        """Adiciona análise de riscos"""
        if not risks:
            return
        
        table = document.add_table(rows=len(risks)+1, cols=4)
        table.style = 'Light Grid'
        
        headers = ['Risco', 'Probabilidade', 'Impacto', 'Mitigação']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, risk in enumerate(risks):
            table.cell(i+1, 0).text = risk.get('risk', '')
            table.cell(i+1, 1).text = risk.get('probability', '')
            table.cell(i+1, 2).text = risk.get('impact', '')
            table.cell(i+1, 3).text = risk.get('mitigation', '')
    
    async def generate_proposal_pdf(self, project_data: Dict) -> str:
        """Gera proposta em formato PDF"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#003366'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        # Título
        story.append(Paragraph(project_data.get('title', 'PROJETO PRONAS/PCD'), title_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Conteúdo
        for section in ['justification', 'objectives', 'methodology']:
            if section in project_data:
                story.append(Paragraph(section.upper(), styles['Heading1']))
                
                if isinstance(project_data[section], dict):
                    for key, value in project_data[section].items():
                        story.append(Paragraph(str(value), styles['BodyText']))
                else:
                    story.append(Paragraph(str(project_data[section]), styles['BodyText']))
                
                story.append(Spacer(1, 0.2*inch))
        
        # Gerar PDF
        doc.build(story)
        buffer.seek(0)
        
        return base64.b64encode(buffer.read()).decode('utf-8')
    
    async def generate_budget_excel(self, project_data: Dict) -> str:
        """Gera planilha de orçamento em Excel"""
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Orçamento"
        
        # Cabeçalho
        ws['A1'] = 'ORÇAMENTO DO PROJETO'
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:E1')
        
        ws['A3'] = 'Categoria'
        ws['B3'] = 'Item'
        ws['C3'] = 'Quantidade'
        ws['D3'] = 'Valor Unitário'
        ws['E3'] = 'Valor Total'
        
        # Estilo do cabeçalho
        for cell in ws[3]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        
        # Dados do orçamento
        budget = project_data.get('budget', {})
        row = 4
        
        for category, items in budget.get('items', {}).items():
            for item in items:
                ws[f'A{row}'] = category.replace('_', ' ').title()
                ws[f'B{row}'] = item.get('description', '')
                ws[f'C{row}'] = 1  # Quantidade padrão
                ws[f'D{row}'] = item.get('value', 0)
                ws[f'E{row}'] = f"=C{row}*D{row}"
                row += 1
        
        # Total
        ws[f'A{row+1}'] = 'TOTAL'
        ws[f'A{row+1}'].font = Font(bold=True)
        ws[f'E{row+1}'] = f"=SUM(E4:E{row-1})"
        ws[f'E{row+1}'].font = Font(bold=True)
        
        # Ajustar largura das colunas
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Salvar
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return base64.b64encode(buffer.read()).decode('utf-8')
    
    async def generate_workplan_pdf(self, project_data: Dict) -> str:
        """Gera plano de trabalho em PDF"""
        return await self.generate_proposal_pdf(project_data)  # Simplificado